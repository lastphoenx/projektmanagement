"""Einheitliches Branding für exportierte DOCX-Dokumente."""

from __future__ import annotations

APP_BRAND_NAME = "Projektmanagement"
FOOTER_TAGLINE = "KI-gestützte Projektplanung"


def export_footer_single_line() -> str:
    return f"Erstellt mit {APP_BRAND_NAME} · {FOOTER_TAGLINE}"


def apply_docx_page_footer(doc) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    text = export_footer_single_line()
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        run = paragraph.add_run(text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
