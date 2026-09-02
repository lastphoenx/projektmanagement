"""Markdown → DOCX Konvertierung (python-docx)."""

from __future__ import annotations

import re
from io import BytesIO

from app.core.export.branding import apply_docx_page_footer


def _add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*\n]+?\*\*|\*[^*\n]+?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def md_to_docx_body(doc, content: str) -> None:
    lines = content.splitlines()
    i = 0
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        if not table_rows:
            return
        cols = len(table_rows[0])
        tbl = doc.add_table(rows=len(table_rows), cols=cols)
        tbl.style = "Table Grid"
        for r_idx, row in enumerate(table_rows):
            for c_idx, cell_text in enumerate(row):
                cell = tbl.cell(r_idx, c_idx)
                cell.paragraphs[0].clear()
                _add_inline_runs(cell.paragraphs[0], cell_text.strip())
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        doc.add_paragraph()
        table_rows.clear()

    while i < len(lines):
        line = lines[i]
        if line.startswith("|"):
            stripped = [c for c in line.split("|") if c != ""]
            if all(re.fullmatch(r"[-: ]+", c.strip()) for c in stripped):
                i += 1
                continue
            table_rows.append(stripped)
            i += 1
            continue
        flush_table()

        if not line.strip():
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=2)
        elif re.match(r"^[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, line[2:].strip())
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\d+\. ", "", line).strip())
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, line)
        i += 1

    flush_table()


def md_to_docx_section(doc, content: str, artifact_label: str) -> None:
    from docx.shared import RGBColor

    heading = doc.add_heading(artifact_label, level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    md_to_docx_body(doc, content)


def build_docx_bytes(
    *,
    title: str,
    sections: list[tuple[str, str]],
    subtitle_lines: list[str] | None = None,
    monospace_sections: set[str] | None = None,
    monospace_labels: set[str] | None = None,
) -> bytes:
    from datetime import date

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    mono = monospace_labels or monospace_sections or set()
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    title_p = doc.add_heading(title, 0)
    title_p.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    for line in subtitle_lines or []:
        p = doc.add_paragraph(line)
        p.runs[0].bold = True

    doc.add_paragraph(f"Erstellt am: {date.today().strftime('%d.%m.%Y')}")
    doc.add_page_break()

    for label, content in sections:
        if not content.strip():
            doc.add_heading(label, level=1)
            p = doc.add_paragraph("(Noch kein Inhalt vorhanden)")
            p.runs[0].italic = True
            doc.add_page_break()
            continue

        if label in mono:
            doc.add_heading(label, level=1)
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            md_to_docx_section(doc, content, label)
        doc.add_page_break()

    apply_docx_page_footer(doc)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
